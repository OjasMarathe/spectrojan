"""Build the final submission .docx by filling the Apart Research template.

We open the template, replace title + abstract in the banner table, drop the
"How to use" info box, then replace all body content with our report.

The markdown source of truth is reports/SpecTrojan_Report.md. This script renders
a subset of Markdown features (headings, bold/italic spans, fenced code, tables,
bulleted/numbered lists, blockquotes) into the docx structure.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from copy import deepcopy
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = Path("/Users/ojas/Downloads/Copy of Apart Research hackathon submission template.docx")
SOURCE_MD = ROOT / "reports" / "SpecTrojan_Report.md"
OUT_PATH = ROOT / "reports" / "SpecTrojan_Submission.docx"

TITLE = "SpecTrojan: Adversarial Specification Validation via Evil Twin Synthesis"
SUBTITLE_LINE = "Track 2 — Specification Validation · The Secure Program Synthesis Hackathon, 2026"
AUTHOR_LINE = "Author: Ojas, ⟨YOUR_AFFILIATION⟩"


# ---------- Markdown parsing -----------------------------------------------

def parse_markdown(text: str) -> list[dict]:
    """Parse the report markdown into a flat list of block descriptors.

    Supported blocks:
      {"kind": "h1"|"h2"|"h3", "text": str}
      {"kind": "p", "spans": [(text, {bold, italic, code})]}
      {"kind": "blockquote", "spans": [...]}
      {"kind": "ul", "items": [list of span-lists]}
      {"kind": "ol", "items": [list of span-lists]}
      {"kind": "code", "text": str, "language": str}
      {"kind": "table", "header": [str...], "rows": [[str...], ...]}
      {"kind": "hr"}
    """
    blocks: list[dict] = []
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append({"kind": "code", "text": "\n".join(code_lines), "language": language})
            continue

        if stripped == "---":
            blocks.append({"kind": "hr"})
            i += 1
            continue

        # Standalone image line:  ![caption](path)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if img_match:
            blocks.append({"kind": "image", "caption": img_match.group(1), "path": img_match.group(2)})
            i += 1
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            blocks.append({"kind": f"h{level}", "text": m.group(2).strip()})
            i += 1
            continue

        # Tables: header | header \n |---|---| \n row | row
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\|?\s*-+", lines[i + 1].strip()):
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip separator
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            blocks.append({"kind": "table", "header": header_cells, "rows": rows})
            continue

        if stripped.startswith("> "):
            spans = parse_spans(stripped[2:])
            blocks.append({"kind": "blockquote", "spans": spans})
            i += 1
            continue

        # bulleted list
        if re.match(r"^[-*]\s+", stripped):
            items: list[list[tuple[str, dict]]] = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(parse_spans(re.sub(r"^[-*]\s+", "", lines[i].strip())))
                i += 1
            blocks.append({"kind": "ul", "items": items})
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(parse_spans(re.sub(r"^\d+\.\s+", "", lines[i].strip())))
                i += 1
            blocks.append({"kind": "ol", "items": items})
            continue

        # paragraph (collect until blank line)
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,3} |```|---|>|[-*]\s+|\d+\.\s+|\|)", lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append({"kind": "p", "spans": parse_spans(" ".join(para_lines))})

    return blocks


def parse_spans(text: str) -> list[tuple[str, dict]]:
    """Tokenize inline markdown into (text, formatting-dict) spans."""
    # Order of precedence: code > bold > italic
    spans: list[tuple[str, dict]] = []

    # Replace markdown links [text](url) → text (url) to keep simple
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)

    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            spans.append((text[pos:m.start()], {}))
        token = m.group(0)
        if token.startswith("`"):
            spans.append((token[1:-1], {"code": True}))
        elif token.startswith("**") or token.startswith("__"):
            spans.append((token[2:-2], {"bold": True}))
        else:
            spans.append((token[1:-1], {"italic": True}))
        pos = m.end()
    if pos < len(text):
        spans.append((text[pos:], {}))
    return spans


# ---------- Rendering into the docx ----------------------------------------

def _set_run_format(run, fmt: dict) -> None:
    if fmt.get("bold"):
        run.bold = True
    if fmt.get("italic"):
        run.italic = True
    if fmt.get("code"):
        run.font.name = "Menlo"
        run.font.size = Pt(9)


def _add_table_borders(table) -> None:
    """Add single-line borders to every cell of `table`."""
    from docx.oxml import OxmlElement
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BBBBBB")
        borders.append(b)
    tblPr.append(borders)


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_spans(para, spans: list[tuple[str, dict]]) -> None:
    for text, fmt in spans:
        run = para.add_run(text)
        _set_run_format(run, fmt)


def remove_element(elem) -> None:
    elem.getparent().remove(elem)


def insert_para_after(prev_elem, doc, style: str | None = None):
    """Insert a new paragraph immediately after prev_elem; return the paragraph object."""
    new_p = doc.paragraphs[0]  # any existing para — we use _body API instead
    from docx.oxml import OxmlElement
    p_xml = OxmlElement("w:p")
    prev_elem.addnext(p_xml)
    # Wrap into python-docx Paragraph
    from docx.text.paragraph import Paragraph
    para = Paragraph(p_xml, doc.paragraphs[0]._parent)
    if style:
        try:
            para.style = doc.styles[style]
        except KeyError:
            pass
    return para


def render(doc: Document, blocks: list[dict]) -> None:
    """Append all blocks to the document body."""
    for block in blocks:
        k = block["kind"]
        if k.startswith("h"):
            level = int(k[1:])
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 2")
            p = doc.add_paragraph(style=style)
            p.add_run(block["text"])
        elif k == "p":
            p = doc.add_paragraph()
            add_spans(p, block["spans"])
        elif k == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run_intro = p.add_run("")
            add_spans(p, block["spans"])
            for run in p.runs:
                run.italic = True
        elif k == "ul":
            for item in block["items"]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.add_run("• ")
                add_spans(p, item)
        elif k == "ol":
            for i_idx, item in enumerate(block["items"], start=1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.add_run(f"{i_idx}. ")
                add_spans(p, item)
        elif k == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(block["text"])
            run.font.name = "Menlo"
            run.font.size = Pt(9)
        elif k == "table":
            header = block["header"]
            rows = block["rows"]
            t = doc.add_table(rows=1 + len(rows), cols=len(header))
            # Manual borders (no theme style required) + header shading.
            _add_table_borders(t)
            for ci, h in enumerate(header):
                cell = t.cell(0, ci)
                cell.text = ""
                _shade_cell(cell, "D9E2F3")
                p = cell.paragraphs[0]
                add_spans(p, parse_spans(h))
                for run in p.runs:
                    run.bold = True
            for ri, row in enumerate(rows, start=1):
                for ci, cv in enumerate(row):
                    if ci >= len(header):
                        continue
                    cell = t.cell(ri, ci)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_spans(p, parse_spans(cv))
            doc.add_paragraph()  # spacer
        elif k == "image":
            img_path = block["path"]
            if not Path(img_path).is_absolute():
                # paths in the markdown are relative to the markdown file's directory
                img_path = str((ROOT / "reports" / img_path).resolve())
            if not Path(img_path).is_file():
                # graceful fallback if the image is missing — emit a placeholder paragraph
                p = doc.add_paragraph()
                run = p.add_run(f"[FIGURE MISSING: {img_path}]")
                run.italic = True
            else:
                doc.add_picture(img_path, width=Inches(6.2))
                # centered caption beneath
                if block["caption"]:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(block["caption"])
                    cap_run.italic = True
                    cap_run.font.size = Pt(10)
        elif k == "hr":
            doc.add_paragraph()  # blank line


def main() -> int:
    if not TEMPLATE.exists():
        print(f"template not found: {TEMPLATE}", file=sys.stderr)
        return 1

    md_text = SOURCE_MD.read_text()
    blocks = parse_markdown(md_text)

    # Strip the top-of-file H1 title, subtitle line, author line, abstract
    # — these go in the banner table, not the body.
    # We assume the structure: H1, then a paragraph (subtitle), then maybe author paragraph,
    # then ---, then "## Abstract", then a paragraph, then ---.
    # Find the abstract paragraph and the index where body content starts.

    abstract_text = ""
    body_start = 0
    for idx, b in enumerate(blocks):
        if b["kind"] == "h2" and b["text"].lower().strip() == "abstract":
            # next paragraph is the abstract
            if idx + 1 < len(blocks) and blocks[idx + 1]["kind"] == "p":
                abstract_text = "".join(s for s, _ in blocks[idx + 1]["spans"])
            # body starts after the next --- or after the abstract paragraph
            for k in range(idx + 2, len(blocks)):
                if blocks[k]["kind"] == "h2":
                    body_start = k
                    break
            break

    if not abstract_text:
        # Fallback: look for inline ## Abstract differently
        abstract_text = "Abstract not found in markdown."

    # Load template
    doc = Document(str(TEMPLATE))

    # ---- Fill title + abstract banner table (table 0) ----
    banner = doc.tables[0]
    title_cell = banner.cell(0, 0)
    title_cell.text = ""
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(18)

    p_sub = title_cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run(SUBTITLE_LINE)
    run.italic = True
    run.font.size = Pt(10)

    p_auth = title_cell.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_auth.add_run(AUTHOR_LINE)
    run.font.size = Pt(10)

    abstract_cell = banner.cell(1, 0)
    abstract_cell.text = ""
    p_abs_h = abstract_cell.paragraphs[0]
    run = p_abs_h.add_run("Abstract")
    run.bold = True
    p_abs = abstract_cell.add_paragraph()
    add_spans(p_abs, parse_spans(abstract_text))

    # ---- Remove the "How to use" info box (table 1) ----
    if len(doc.tables) > 1:
        info_table = doc.tables[1]
        info_table._element.getparent().remove(info_table._element)

    # ---- Remove ALL existing body paragraphs (guidance text + section headings) ----
    # Strategy: iterate through doc.element.body and remove every <w:p> that comes AFTER
    # the banner table. Then we'll append our content fresh.
    body = doc.element.body
    # Locate the banner table element to anchor our deletions
    banner_elem = banner._element

    # Remove all <w:p> and remaining <w:tbl> elements after the banner.
    to_remove = []
    seen_banner = False
    for child in list(body):
        if child is banner_elem:
            seen_banner = True
            continue
        if not seen_banner:
            continue
        # keep the final sectPr
        if child.tag == qn("w:sectPr"):
            continue
        to_remove.append(child)
    for child in to_remove:
        body.remove(child)

    # ---- Append our body content ----
    # Drop initial blocks (title/abstract front matter) — only render from body_start onward
    body_blocks = blocks[body_start:]
    render(doc, body_blocks)

    doc.save(str(OUT_PATH))
    print(f"wrote {OUT_PATH}")
    print(f"file size: {OUT_PATH.stat().st_size:,} bytes")
    return 0


if __name__ == "__main__":
    sys.exit(main())
